"""Convert draft_prateek_UPDATED.md to draft_prateek_UPDATED.docx (requires python-docx)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
IMG_LINE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
MD_PATH = ROOT / "draft_prateek_UPDATED.md"
OUT_PATH = ROOT / "draft_prateek_UPDATED.docx"


def add_formatted_paragraph(doc: Document, text: str) -> None:
    """Paragraph with simple **bold** segments."""
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def add_code_block(doc: Document, lines_buf: list[str]) -> None:
    """Fenced code: monospace, one paragraph line each (preserves alignment in Word)."""
    if not lines_buf:
        return
    for raw in lines_buf:
        line_text = raw.rstrip("\r\n")
        p = doc.add_paragraph()
        run = p.add_run(line_text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    doc.add_paragraph()


def main() -> None:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    title_done = False
    while i < len(lines):
        line = lines[i].rstrip()
        if line.strip() == "---":
            i += 1
            continue
        if line.startswith("# ") and not title_done:
            t = doc.add_heading(line[2:].strip(), level=0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_done = True
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in table_lines if "---" not in r]
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        tbl.cell(ri, ci).text = cell
            continue
        if line.strip().startswith("```"):
            buf: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            add_code_block(doc, buf)
            continue
        img_match = IMG_LINE.match(line.strip())
        if img_match:
            alt, rel = img_match.groups()
            parts = [p for p in rel.replace("\\", "/").split("/") if p and p != "."]
            img_path = ROOT.joinpath(*parts).resolve()
            cap = doc.add_paragraph()
            cap.add_run(alt).italic = True
            if img_path.is_file():
                doc.add_picture(str(img_path), width=Inches(6.3))
            else:
                doc.add_paragraph(f"[Missing image file: {rel}]")
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        add_formatted_paragraph(doc, line.strip())
        i += 1

    try:
        doc.save(OUT_PATH)
        print(f"Wrote {OUT_PATH}")
    except PermissionError:
        alt = ROOT / "draft_prateek_UPDATED_SCREENSHOTS.docx"
        doc.save(alt)
        print(f"Primary locked; wrote {alt}")


if __name__ == "__main__":
    main()
